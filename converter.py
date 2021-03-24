import re
import mammoth
from atom_weight import atom_weight
from bs4 import BeautifulSoup
import docx


forbidden_words = ['Макс.', 'Среднее', 'Станд. отклонение', 'стат.', 'отклонение', 'Станд.', 'Мин.']
forbidden_values = [ 'Да', 'Итог', '100.00', 'В стат.']
tables = []
converted_values = []

with open("138.docx", "rb") as docx_file:
    result = mammoth.convert_to_html(docx_file)
    html = result.value # The generated HTML
    messages = result.messages # Any messages, such as warnings during conversion


def extract_tables(row_html):
    tables_list = re.findall(r'<table>((.|\n)*?)<\/table>', row_html)
    return tables_list


def extract_rows(row_table):
    return re.findall(r'<tr>((.|\n)*?)<\/tr>', row_table)


def clean_tags(row_html):
    return re.split('<.*?>', row_html)


def add_zero_value(row):
    new_row = ""
    soup = BeautifulSoup(row, 'html.parser')
    tags = soup.find_all("td")
    for tag in tags:
        tag = str(tag)
        if tag == "<td></td>":
            tag = "<td>0</td>"
            new_row += tag
        else:
            new_row += tag
    return new_row


def clean_row(row):
    new_row = []
    for item in row:
        if not item in forbidden_words:
            new_row.append(item)
    return new_row


def remove_spaces(row):
    new_row = []
    for item in row:
        if item != '':
            new_row.append(item)
    return new_row


def remove_emtpy_row(row):
    if row[0] == '0':
        return
    return row


def remove_forbidden_values(row):
    new_row = []
    for item in row:
        if not item in forbidden_values:
            new_row.append(item)
    return new_row


html_tuple = extract_tables(html)


for table in html_tuple:
    tables.append([])
    for row in extract_rows(table[0]):
        string_without_spaces = add_zero_value(row[0])
        cleaned_row = clean_tags(string_without_spaces)
        cleaned_row = remove_spaces(cleaned_row)
        cleaned_row = remove_emtpy_row(cleaned_row)
        # print("cleaned_row", cleaned_row)
        if cleaned_row != None and cleaned_row[0] not in forbidden_words:
            cleaned_row = remove_forbidden_values(cleaned_row)
            tables[-1].append(cleaned_row)


spectrus = []
atomic_weight = []
temp_table = []

for table in tables:
    new_table = []
    header = table[0]
    new_table.append(header)
    table_length = len(table)

    current_position = 1
    current_row = 1

    for cycle in range(table_length-1):
        new_table.append([])
        new_table[current_row].append(table[current_row][0])
        current_row += 1
    current_row = 1

    for element in table[0][1:]:
        for atomic in atom_weight:
            if element == atomic[0]:
                while current_row < table_length:
                    # if element == "O":
                        # print("Oxigen", table[current_row][0], table[current_row][current_position], atomic[1])
                        # print("value", (float(table[current_row][current_position]) / atomic[1]))
                    new_table[current_row].append(float(table[current_row][current_position])/atomic[1])
                    current_row += 1
        current_position += 1
        current_row = 1
    temp_table.append(new_table)

final_table = []

for table in temp_table:
    new_table = []
    row_sum = None
    header = table[0]
    new_table.append(header)
    table_length = len(table)

    current_position = 1
    current_row = 1
    for cycle in range(table_length-1):
        new_table.append([])
        new_table[current_row].append(table[current_row][0])

        # print(table[current_row][0], row_sum)
        current_row += 1
    current_row = 1
    for element in table[0][1:]:
        while current_row < table_length:
            row_value = [float(item) for item in table[current_row][1:]]
            row_sum = sum(row_value)
            # print(table[current_row][0], "position", current_row, current_position, table[current_row][current_position])
            value = round(float(table[current_row][current_position]) / row_sum * 100, 2)
            # print(row_sum)
            new_table[current_row].append(value)
            current_row += 1
        current_position += 1

        current_row = 1
    final_table.append(new_table)
print(final_table)


source_document = docx.Document("138.docx")
target_document = docx.Document()

p = source_document.add_paragraph()
r = p.add_run()
r.add_text('Good Morning every body,This is my ')
r.add_picture('12345.jpg')
r.add_text(' do you like it?')
source_document.save("test.docx")
